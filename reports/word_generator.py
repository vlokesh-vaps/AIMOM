"""Word (DOCX) MoM generator using python-docx."""

import time
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from config.settings import COMPANY_NAME, COMPANY_THEME_COLOR
from utils.logger import get_logger

logger = get_logger(__name__)


def _hex_to_rgb(hex_color: str) -> RGBColor:
    value = hex_color.lstrip("#")
    if len(value) != 6:
        value = "1E3A8A"
    return RGBColor(int(value[0:2], 16), int(value[2:4], 16), int(value[4:6], 16))


def _set_cell_shading(cell, hex_color: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), hex_color.lstrip("#"))
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


class WordGenerator:
    """Generate a detailed DOCX Minutes of Meeting report."""

    def __init__(self) -> None:
        self._output_dir = Path(__file__).resolve().parent / "word"
        self._output_dir.mkdir(parents=True, exist_ok=True)

    def generate(self, summary_data: dict, filename: str) -> Path:
        start_time = time.time()
        output_path = self._output_dir / filename

        doc = Document()
        theme_color = _hex_to_rgb(COMPANY_THEME_COLOR)
        self._configure_document(doc)

        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("MINUTES OF MEETING")
        run.bold = True
        run.font.size = Pt(16)
        run.font.color.rgb = theme_color

        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run(summary_data.get("meeting_title", "Meeting"))
        run.bold = True
        run.font.size = Pt(13)

        meta_table = doc.add_table(rows=6, cols=2)
        meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        meta_rows = [
            ("Date", summary_data.get("meeting_date", "") or self._format_datetime(summary_data.get("generated_at", ""))),
            ("Chaired By", summary_data.get("chaired_by", "") or "Not specified"),
            ("Meeting Type", summary_data.get("meeting_type", "General")),
            ("Organization", summary_data.get("organization", "") or COMPANY_NAME),
            ("Attendees", self._join_people(summary_data.get("attendees") or summary_data.get("participants") or [])),
            ("Absents", summary_data.get("absents", "") or "Nil"),
        ]
        self._fill_meta_table(meta_table, meta_rows, theme_color)

        doc.add_paragraph("")

        table = doc.add_table(rows=1, cols=6)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"

        headers = [
            "S.No",
            "Agenda",
            "Discussion",
            "Action Item",
            "Assigned",
            "Target Date",
        ]
        header_row = table.rows[0]
        for idx, header in enumerate(headers):
            cell = header_row.cells[idx]
            cell.text = header
            _set_cell_shading(cell, COMPANY_THEME_COLOR)
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        for serial, agenda, summary_text, action_text, owner_text, date_text in self._build_rows(summary_data):
            row = table.add_row().cells
            row[0].text = str(serial)
            row[1].text = agenda
            row[2].text = summary_text
            row[3].text = action_text
            row[4].text = owner_text
            row[5].text = date_text
            for cell in row:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)

        widths = [Cm(1.2), Cm(4.2), Cm(7.4), Cm(5.8), Cm(3.4), Cm(3.2)]
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = width

        doc.add_paragraph("")
        note = doc.add_paragraph(
            "Note: This report was generated from the meeting transcript. Please review before circulation."
        )
        note.runs[0].italic = True
        note.runs[0].font.size = Pt(8)

        doc.save(str(output_path))
        logger.info("Word generation complete. Path=%s, Duration=%.2fs", output_path, time.time() - start_time)
        return output_path

    @staticmethod
    def _configure_document(doc: Document) -> None:
        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(10)

    @staticmethod
    def _fill_meta_table(table, rows: list[tuple[str, str]], theme_color: RGBColor) -> None:
        for index, (label, value) in enumerate(rows):
            label_cell = table.cell(index, 0)
            value_cell = table.cell(index, 1)
            label_cell.text = label
            value_cell.text = value
            _set_cell_shading(label_cell, "F1F5F9")
            for paragraph in label_cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.color.rgb = theme_color
                    run.font.size = Pt(9)
            for paragraph in value_cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)

    @staticmethod
    def _build_rows(summary_data: dict) -> list[tuple[int, str, str, str, str, str]]:
        discussion_points = summary_data.get("discussion_points", [])
        action_items = summary_data.get("action_items", [])
        actions_by_agenda: dict[str, list] = {}
        for item in action_items:
            agenda = WordGenerator._get(item, "agenda_item", "").strip() or "Off Agenda Discussion"
            actions_by_agenda.setdefault(agenda, []).append(item)

        rows: list[tuple[int, str, str, str, str, str]] = []
        used_action_ids: set[int] = set()

        for idx, dp in enumerate(discussion_points, start=1):
            agenda = WordGenerator._get(dp, "agenda_item", "Off Agenda Discussion").strip() or "Off Agenda Discussion"
            matched = actions_by_agenda.get(agenda, [])
            agenda_label = agenda
            action_text = WordGenerator._join_lines(WordGenerator._get(ai, "task", "") for ai in matched)
            if not action_text:
                fallback_task = WordGenerator._get(dp, "task", "")
                action_text = "" if fallback_task == "No Action Item" else fallback_task

            summary_parts = [WordGenerator._get(dp, "point", ""), WordGenerator._get(dp, "detailed_summary", "")]
            decision = WordGenerator._get(dp, "decision", "")
            if decision and decision != "No Decision Taken":
                summary_parts.append(f"Decision: {decision}")
            summary_text = WordGenerator._join_non_empty(summary_parts)

            owner_text = WordGenerator._join_lines(WordGenerator._get(ai, "owner", "") for ai in matched)
            if not owner_text:
                owner_text = WordGenerator._get(dp, "assigned_to", "Not Specified")

            date_text = WordGenerator._join_lines(WordGenerator._get(ai, "target_date", "") for ai in matched)
            if not date_text:
                date_text = WordGenerator._get(dp, "deadline", "Not Specified")

            rows.append((idx, agenda_label, summary_text, action_text or "No Action Item", owner_text, date_text))
            for item in matched:
                used_action_ids.add(id(item))

        next_index = len(rows) + 1
        for item in action_items:
            if id(item) in used_action_ids:
                continue
            agenda = WordGenerator._get(item, "agenda_item", "").strip() or "Off Agenda Discussion"
            rows.append((
                next_index,
                agenda,
                "",
                WordGenerator._get(item, "task", ""),
                WordGenerator._get(item, "owner", ""),
                WordGenerator._get(item, "target_date", ""),
            ))
            next_index += 1

        return rows

    @staticmethod
    def _join_people(values: list[str]) -> str:
        cleaned = [str(value).strip() for value in values if str(value).strip()]
        return " | ".join(cleaned) if cleaned else "Not provided"

    @staticmethod
    def _join_lines(values) -> str:
        seen: list[str] = []
        for value in values:
            text = str(value or "").strip()
            if not text or text in ("Not Specified", "-"):
                continue
            if text not in seen:
                seen.append(text)
        return "\n".join(seen)

    @staticmethod
    def _join_non_empty(values: list[str]) -> str:
        seen: list[str] = []
        for value in values:
            text = str(value or "").strip()
            if not text:
                continue
            if text not in seen:
                seen.append(text)
        return "\n".join(seen)

    @staticmethod
    def _get(item, key: str, default: str = "") -> str:
        if isinstance(item, dict):
            return item.get(key, default)
        return getattr(item, key, default)

    @staticmethod
    def _format_datetime(value: str) -> str:
        if not value:
            return ""
        try:
            parsed = datetime.fromisoformat(str(value).replace("Z", "+00:00"))
        except ValueError:
            return str(value)
        return parsed.strftime("%d %B %Y")
